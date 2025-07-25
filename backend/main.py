#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
官方AI写作系统 - 主应用入口
集成第一阶段优化：环境变量、安全配置、输入验证、异常处理、代码去重
"""

# 初始化SQLite（消除重复代码）
try:
    from utils.sqlite_init import init_sqlite
    init_sqlite()
except ImportError as e:
    print(f"警告：SQLite初始化失败: {e}")
    # 如果导入失败，使用基础的SQLite初始化
    import sys
    import os
    from pathlib import Path
    
    venv_path = Path(__file__).parent / "venv"
    site_packages = list(venv_path.glob("lib/python*/site-packages"))
    if site_packages:
        pysqlite3_path = site_packages[0] / "pysqlite3"
        if pysqlite3_path.exists():
            sys.path.insert(0, str(pysqlite3_path))
    
    try:
        import pysqlite3
        import sqlite3
        sys.modules['sqlite3'] = pysqlite3
        print("✓ 已切换到pysqlite3")
    except ImportError:
        print("⚠ pysqlite3导入失败，使用系统sqlite3")
    
    import sqlite3
    print(f"当前使用的SQLite版本: {sqlite3.sqlite_version}")

# 导入第一阶段优化的模块
try:
    from config.security import init_security_config, security_config
    from utils.validators import InputValidator, FileUploadValidator, ContentValidator
    from utils.error_handler import register_error_handlers, create_validation_error, create_file_upload_error
    from utils.logger import get_app_logger, setup_logging
except ImportError as e:
    print(f"警告：无法导入优化模块: {e}")
    # 使用基础日志配置
    import logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    logger = logging.getLogger('backend.app')
else:
    # 初始化优化模块
    try:
        setup_logging()
        logger = get_app_logger()
        init_security_config()
        logger.info("第一阶段优化模块初始化成功")
    except Exception as e:
        print(f"警告：优化模块初始化失败: {e}")
        import logging
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger('backend.app')

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import mysql.connector
import os
import json
from werkzeug.utils import secure_filename
import requests
from datetime import datetime
import tempfile
import docx
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 导入知识库相关模块
try:
    from routes.knowledge_base_simple import knowledge_base_simple_bp as knowledge_base_bp
except ImportError:
    # 如果导入失败，尝试使用相对导入
    try:
        from backend.routes.knowledge_base_simple import knowledge_base_simple_bp as knowledge_base_bp
    except ImportError:
        logger.warning("无法导入知识库模块")
        knowledge_base_bp = None

# 导入RAG生成相关模块
try:
    from routes.rag_generation import rag_generation_bp
except ImportError:
    # 如果导入失败，尝试使用相对导入
    try:
        from backend.routes.rag_generation import rag_generation_bp
    except ImportError:
        logger.warning("无法导入RAG生成模块")
        rag_generation_bp = None

# 导入AI操作相关模块
try:
    from routes.ai_operations import ai_operations_bp
except ImportError:
    # 如果导入失败，尝试使用相对导入
    try:
        from backend.routes.ai_operations import ai_operations_bp
    except ImportError:
        logger.warning("无法导入AI操作模块")
        ai_operations_bp = None

# 导入配置（使用新的安全配置模块）
try:
    from config.security import security_config
    DB_CONFIG = security_config.DB_CONFIG
    DEEPSEEK_API_URL = security_config.DEEPSEEK_API_URL
    DEEPSEEK_API_KEY = security_config.DEEPSEEK_API_KEY
    UPLOAD_FOLDER = security_config.get_safe_config_dict()['file']['upload_folder']
    ALLOWED_EXTENSIONS = security_config.get_safe_config_dict()['file']['allowed_extensions']
except ImportError:
    # 如果导入失败，使用默认配置
    logger.warning("无法导入安全配置，使用默认配置")
    DB_CONFIG = {
        'host': '47.118.250.53',
        'port': 3306,
        'database': 'official_doc',
        'user': 'official_doc',
        'password': 'your-password',
        'charset': 'utf8mb4'
    }
    DEEPSEEK_API_URL = 'https://api.deepseek.com/v1/chat/completions'
    DEEPSEEK_API_KEY = 'your-api-key'
    UPLOAD_FOLDER = 'temp'
    ALLOWED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt', '.md']

app = Flask(__name__)

# 使用安全配置
try:
    app.config['MAX_CONTENT_LENGTH'] = security_config.get_safe_config_dict()['file']['max_file_size']
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
    app.config['SECRET_KEY'] = security_config.JWT_SECRET_KEY
    app.debug = security_config.get_safe_config_dict()['debug']
except Exception as e:
    logger.warning(f"使用安全配置失败，使用默认配置: {e}")
    app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
    app.config['SECRET_KEY'] = 'default-secret-key'
    app.debug = True

# 配置CORS，允许特定域名访问
CORS(app, origins=[
    'http://localhost:8081',
    'http://127.0.0.1:8081',
    'http://115.190.152.96:8081',
    'http://chenxiaoshivivid.com.cn:8081',
    'http://localhost:8005',
    'http://121.36.205.70:8005',
    'http://chenxiaoshivivid.com.cn:8005'
])

# 注册全局错误处理器
try:
    register_error_handlers(app)
    logger.info("全局错误处理器注册成功")
except Exception as e:
    logger.warning(f"全局错误处理器注册失败: {e}")

# 注册知识库蓝图
if knowledge_base_bp:
    app.register_blueprint(knowledge_base_bp, url_prefix='/api/knowledge')
    logger.info("知识库蓝图注册成功")
else:
    logger.warning("知识库蓝图未注册")

# 注册RAG生成蓝图
if rag_generation_bp:
    app.register_blueprint(rag_generation_bp, url_prefix='/api/rag')
    logger.info("RAG生成蓝图注册成功")
else:
    logger.warning("RAG生成蓝图未注册")

# 注册AI操作蓝图
if ai_operations_bp:
    app.register_blueprint(ai_operations_bp, url_prefix='/api/ai')
    logger.info("AI操作蓝图注册成功")
else:
    logger.warning("AI操作蓝图未注册")

# 简化请求日志中间件
@app.before_request
def log_request_info():
    logger.info(f'请求: {request.method} {request.path}')

# 简化响应日志中间件
@app.after_request
def log_response_info(response):
    logger.info(f'响应: {response.status_code}')
    return response

# 错误处理中间件
@app.errorhandler(400)
def handle_bad_request(error):
    """处理400错误，特别是SSL/TLS协议错误"""
    if 'Bad request version' in str(error):
        logger.warning(f"检测到SSL/TLS协议错误，可能是HTTPS请求发送到HTTP服务器: {error}")
        return jsonify({
            'error': '协议错误',
            'message': '请使用HTTP协议访问此服务',
            'status': 'protocol_error'
        }), 400
    return jsonify({'error': '请求错误'}), 400

@app.errorhandler(404)
def handle_not_found(error):
    """处理404错误"""
    logger.warning(f"请求的路径不存在: {request.path}")
    return jsonify({
        'error': '路径不存在',
        'message': f'请求的路径 {request.path} 不存在',
        'available_paths': ['/api/templates', '/api/generate', '/api/upload', '/signin']
    }), 404

# 确保上传文件夹存在
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    """检查文件是否允许上传（使用新的验证器）"""
    try:
        return InputValidator.validate_filename(filename)
    except Exception:
        # 如果验证器不可用，使用原有逻辑
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_db_connection():
    """获取数据库连接"""
    try:
        # 尝试使用安全配置
        try:
            from config.security import security_config
            return mysql.connector.connect(**security_config.DB_CONFIG)
        except ImportError:
            # 如果无法导入安全配置，使用默认配置
            return mysql.connector.connect(**DB_CONFIG)
    except mysql.connector.Error as err:
        logger.error(f"数据库连接错误: {err}")
        return None

def replace_text_in_paragraphs(paragraphs, replacements):
    """支持跨run的文本替换，并处理段落格式"""
    for paragraph in paragraphs:
        # 合并所有run的文本
        full_text = ''.join(run.text for run in paragraph.runs)
        replaced = False
        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)
                replaced = True
        if replaced and paragraph.runs:
            # 清空原有runs
            for run in paragraph.runs:
                run.text = ''
            # 只用第一个run写入新内容
            paragraph.runs[0].text = full_text

def apply_document_formatting(doc, content):
    """应用公文格式要求"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    
    # 设置段落格式
    for paragraph in doc.paragraphs:
        # 设置行距为30磅
        paragraph.paragraph_format.line_spacing = Pt(30)
        
        # 根据内容设置字体格式
        for run in paragraph.runs:
            text = run.text.strip()
            
            # 一级标题：一、二、三、... - 三号黑体字
            if text and text[0] in '一二三四五六七八九十' and text.endswith('、'):
                run.font.name = '黑体'
                run.font.size = Pt(16)  # 三号字约16磅
                run.font.bold = True
            
            # 二级标题：（一）（二）（三）... - 三号楷体_GB2312字
            elif text and text.startswith('（') and text.endswith('）') and len(text) == 4:
                run.font.name = '楷体_GB2312'
                run.font.size = Pt(16)
            
            # 三级标题：1. 2. 3. ... - 三号仿宋_GB2312字，加粗
            elif text and text[0].isdigit() and text.endswith('.'):
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(16)
                run.font.bold = True
            
            # 正文：三号仿宋_GB2312字
            else:
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(16)


def replace_text_in_document(doc, replacements):
    # 段落
    replace_text_in_paragraphs(doc.paragraphs, replacements)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)

# 在现有路由之前添加signin路由
@app.route('/signin', methods=['GET', 'POST'])
def signin():
    """处理登录请求"""
    logger.info(f"请求: {request.method} /signin")
    
    if request.method == 'GET':
        # GET请求返回登录页面或重定向到前端
        return jsonify({
            'message': '请使用POST方法进行登录',
            'status': 'redirect',
            'frontend_url': 'http://localhost:8005'
        }), 200
    
    elif request.method == 'POST':
        # POST请求处理登录逻辑
        try:
            data = request.get_json()
            username = data.get('username', '')
            password = data.get('password', '')
            
            # 这里可以添加实际的用户验证逻辑
            # 目前返回成功响应
            logger.info(f"登录请求: 用户 {username}")
            
            return jsonify({
                'success': True,
                'message': '登录成功',
                'user': username
            }), 200
            
        except Exception as e:
            logger.error(f"登录处理错误: {e}")
            return jsonify({
                'success': False,
                'message': '登录失败'
            }), 400

@app.route('/api/templates', methods=['GET'])
def get_templates():
    """获取公文类型列表"""
    logger.info("获取模板列表")
    conn = get_db_connection()
    if not conn:
        # 如果数据库连接失败，返回本地定义的模板
        try:
            from config import DOCUMENT_TYPES
        except ImportError:
            from backend.config import DOCUMENT_TYPES
        templates = [{'id': dt[0], 'name': dt[1], 'description': dt[2]} for dt in DOCUMENT_TYPES]
        return jsonify({'data': templates})
    
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM document_types ORDER BY id")
    templates = cursor.fetchall()
    cursor.close()
    conn.close()
    
    return jsonify({'data': templates})

@app.route('/api/generate-title', methods=['POST'])
def generate_title():
    """从内容生成标题"""
    data = request.json
    content = data.get('content', '')
    
    if not content:
        return jsonify({'success': False, 'message': '内容不能为空'})
    
    try:
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {
                        "role": "system", 
                        "content": "你是一个专业的公文写作助手，请根据提供的公文内容生成一个简洁、准确的标题。标题应该符合公文格式规范。"
                    },
                    {
                        "role": "user", 
                        "content": f"请根据以下公文内容生成标题：\n\n{content}"
                    }
                ]
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            title = result['choices'][0]['message']['content'].strip()
            logger.info(f"生成标题成功: {title}")
            return jsonify({'success': True, 'title': title})
        else:
            logger.error(f"API调用失败: {response.status_code}")
            return jsonify({'success': False, 'message': 'API调用失败'})
            
    except Exception as e:
        logger.error(f"生成标题错误: {e}")
        return jsonify({'success': False, 'message': '生成标题失败，请稍后再试'})

@app.route('/api/generate-content', methods=['POST'])
def generate_content():
    """从主题生成内容（支持参考文件）"""
    data = request.json
    topic = data.get('topic', '')
    document_type = data.get('document_type', '')
    reference_files = data.get('reference_files', [])
    use_reference_files = data.get('use_reference_files', True)
    
    if not topic:
        return jsonify({'success': False, 'message': '主题不能为空'})
    
    try:
        # 构建系统提示词
        system_prompt = f"你是一个专业的公文写作助手，请根据提供的主题生成一篇符合{document_type}格式规范的公文内容。内容应该结构清晰、语言规范、符合公文写作要求。请直接返回正文内容，不要包含标题。"
        
        # 构建用户提示词
        user_prompt = f"请根据以下主题生成{document_type}内容：\n\n{topic}"
        
        # 如果有参考文件且用户选择使用参考文件
        if reference_files and use_reference_files:
            try:
                # 从知识库获取参考文件内容
                reference_content = ""
                for file_id in reference_files:
                    # 查询文件信息
                    conn = get_db_connection()
                    if conn:
                        cursor = conn.cursor(dictionary=True)
                        cursor.execute("SELECT * FROM knowledge_files WHERE id = %s", (file_id,))
                        file_info = cursor.fetchone()
                        cursor.close()
                        conn.close()
                        
                        if file_info and file_info.get('metadata'):
                            metadata = json.loads(file_info['metadata'])
                            if 'content' in metadata:
                                reference_content += f"\n\n参考文件内容：\n{metadata['content'][:1000]}..."  # 限制长度
                
                if reference_content:
                    user_prompt += f"\n\n参考文件内容：{reference_content}"
                    system_prompt += "请结合参考文件的内容，确保生成的内容与参考文件保持一致性和相关性。"
                
            except Exception as e:
                logger.warning(f"获取参考文件内容失败: {e}")
                # 即使参考文件获取失败，也继续生成内容
        
        # 调用AI接口
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {
                        "role": "system", 
                        "content": system_prompt
                    },
                    {
                        "role": "user", 
                        "content": user_prompt
                    }
                ],
                "max_tokens": 2000,
                "temperature": 0.7
            },
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content'].strip()
            logger.info(f"生成内容成功，长度: {len(content)}")
            return jsonify({'success': True, 'content': content})
        else:
            logger.error(f"API调用失败: {response.status_code}")
            return jsonify({'success': False, 'message': 'API调用失败'})
            
    except Exception as e:
        logger.error(f"生成内容错误: {e}")
        return jsonify({'success': False, 'message': '生成内容失败，请稍后再试'})

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文件并解析内容（集成第一阶段优化）"""
    logger.info("接收到文件上传请求")
    
    try:
        # 使用新的验证器
        if 'file' not in request.files:
            raise create_file_upload_error("没有文件在请求中")
        
        file = request.files['file']
        if file.filename == '':
            raise create_file_upload_error("没有选择文件")
        
        logger.info(f"上传的文件: {file.filename}")
        
        # 使用新的文件验证器
        try:
            is_valid, message = FileUploadValidator.validate_upload(file)
            if not is_valid:
                raise create_file_upload_error(message, file.filename)
        except Exception:
            # 如果新验证器不可用，使用原有验证
            if not allowed_file(file.filename):
                raise create_file_upload_error("不支持的文件类型", file.filename)
        
        # 保存并验证文件
        try:
            is_saved, message, file_path = FileUploadValidator.save_and_validate_file(file, UPLOAD_FOLDER)
            if not is_saved:
                raise create_file_upload_error(message, file.filename)
        except Exception:
            # 如果新保存器不可用，使用原有逻辑
            filename = secure_filename(file.filename)
            if not os.path.exists(UPLOAD_FOLDER):
                os.makedirs(UPLOAD_FOLDER)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
        
        # 解析文件内容
        content = ""
        try:
            if file_path.endswith('.md'):
                logger.info("解析Markdown文件")
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_path.endswith('.docx'):
                logger.info("解析Word文件")
                doc = docx.Document(file_path)
                content = '\n'.join([para.text for para in doc.paragraphs])
            elif file_path.endswith('.txt'):
                logger.info("解析文本文件")
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            logger.info(f"文件解析成功，内容长度: {len(content)}")
        except Exception as parse_error:
            logger.error(f"文件解析错误: {parse_error}")
            raise create_file_upload_error(f"文件解析失败: {str(parse_error)}", file.filename)
        finally:
            # 删除临时文件
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                logger.error(f"删除临时文件失败: {e}")
        
        return jsonify({'success': True, 'content': content})
        
    except Exception as e:
        logger.error(f"文件上传错误: {e}")
        if hasattr(e, 'status_code'):
            return jsonify({'success': False, 'message': str(e)}), e.status_code
        else:
            return jsonify({'success': False, 'message': f'文件上传失败: {str(e)}'})

@app.route('/api/generate', methods=['POST'])
def generate_document():
    """生成公文"""
    logger.info("接收到生成公文请求")
    
    try:
        data = request.json
        template_type = data.get('template_type', '')
        metadata = data.get('metadata', {})
        content = data.get('content', '')
        
        logger.info(f"模板类型: {template_type}, 标题: {metadata.get('title', '')}")
        
        # 保存到数据库
        document_id = None
        try:
            conn = get_db_connection()
            if conn:
                cursor = conn.cursor()
                try:
                    cursor.execute("""
                        INSERT INTO generated_documents 
                        (document_type_id, title, content, metadata) 
                        VALUES (%s, %s, %s, %s)
                    """, (
                        template_type,
                        metadata.get('title', ''),
                        content,
                        json.dumps(metadata)
                    ))
                    document_id = cursor.lastrowid
                    conn.commit()
                    logger.info(f"数据已保存到数据库，ID: {document_id}")
                except Exception as db_error:
                    logger.error(f"数据库保存错误: {db_error}")
                finally:
                    cursor.close()
                    conn.close()
        except Exception as conn_error:
            logger.error(f"数据库连接错误: {conn_error}")
        
        # 使用docx模板生成Word文档
        try:
            logger.info("开始使用模板生成Word文档")
            
            # 获取模板文件路径
            template_dir = os.path.join(os.path.dirname(__file__), '..', 'frontend', 'templates')
            template_file = os.path.join(template_dir, f'{template_type}.docx')
            
            if not os.path.exists(template_file):
                logger.error(f"模板文件不存在: {template_file}")
                return jsonify({'success': False, 'message': '模板文件不存在'})
            
            # 加载模板文档
            doc = Document(template_file)
            
            # 准备替换内容 - 支持现有模板格式和新的{{}}格式
            replacements = {
                # 版头字段 - 支持现有格式
                '××★1年': metadata.get('year', '2025年'),
                '特急': metadata.get('urgencyLevel', ''),
                '机关代字(20××)×号': f"{metadata.get('senderCode', '机关代字')}({metadata.get('year', '2025')}){metadata.get('serialNumber', '')}号",
                '签发人:姓名一姓名二': f"签发人:{metadata.get('senderSignature', '')}",
                
                # 主体字段 - 支持现有格式
                '××××关于××××XXXXXXXX×××的报告': metadata.get('title', ''),
                '××××关于××××XXXXXXXX×××的纪要': metadata.get('title', ''),
                '主送机关:': f"主送机关:{metadata.get('recipient', '')}",
                '正文内容': content,
                
                # 署名字段 - 支持现有格式
                '机关短署名': metadata.get('senderSignature', ''),
                '20××年×月×日': metadata.get('date', ''),
                '(附注内容)': f"({metadata.get('notes', '')})",
                
                # 版记字段 - 支持现有格式
                '抄送:抄送机关1,抄送机关2,抄送机关3,抄送机关4,抄送机关5。': f"抄送:{metadata.get('copyTo', '')}",
                
                # 印发字段 - 支持现有格式
                '印发机关': metadata.get('printingOrg', ''),
                '20××年×月×日印发': f"{metadata.get('printingDate', '')}印发",
                
                # 同时支持新的{{}}格式
                '{{copyNumber}}': metadata.get('copyNumber', '000001'),
                '{{securityLevel}}': metadata.get('securityLevel', '一般'),
                '{{securityPeriod}}': metadata.get('securityPeriod', '1年'),
                '{{urgencyLevel}}': metadata.get('urgencyLevel', '特急'),
                '{{sender}}': metadata.get('sender', '省委宣传部'),
                '{{senderSymbol}}': metadata.get('senderSymbol', '文件'),
                '{{senderCode}}': metadata.get('senderCode', '机关代字'),
                '{{year}}': metadata.get('year', '2025'),
                '{{serialNumber}}': metadata.get('serialNumber', '1'),
                '{{senderSignature}}': metadata.get('senderSignature', '姓名1'),
                '{{title}}': metadata.get('title', ''),
                '{{recipient}}': metadata.get('recipient', ''),
                '{{content}}': content,
                '{{date}}': metadata.get('date', ''),
                '{{notes}}': metadata.get('notes', ''),
                '{{copyTo}}': metadata.get('copyTo', '杭州市委宣传部'),
                '{{printingOrg}}': metadata.get('printingOrg', '印发机关'),
                '{{printingDate}}': metadata.get('printingDate', '')
            }
            
            # 执行文本替换
            replace_text_in_document(doc, replacements)
            
            # 应用公文格式要求
            apply_document_formatting(doc, content)
            
            # 保存文档
            temp_dir = tempfile.gettempdir()
            os.makedirs(temp_dir, exist_ok=True)
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_file_path = temp_file.name
            temp_file.close()
            
            doc.save(temp_file_path)
            
            download_url = f'/api/download/{os.path.basename(temp_file_path)}'
            logger.info(f"文档生成成功: {download_url}")
            
            return jsonify({
                'success': True,
                'document_id': document_id,
                'download_url': download_url
            })
            
        except Exception as doc_error:
            logger.error(f"文档生成错误: {doc_error}")
            return jsonify({'success': False, 'message': f'文档生成失败: {str(doc_error)}'})
            
    except Exception as e:
        logger.error(f"处理请求错误: {e}")
        return jsonify({'success': False, 'message': f'处理请求失败: {str(e)}'})

@app.route('/api/preview/<filename>')
def preview_file(filename):
    """预览生成的文档"""
    logger.info(f"请求预览文件: {filename}")
    
    try:
        file_path = os.path.join(tempfile.gettempdir(), filename)
        
        if os.path.exists(file_path):
            logger.info(f"文件存在，准备预览: {file_path}")
            return send_file(file_path, as_attachment=False, download_name='公文.docx')
        else:
            logger.error(f"文件不存在: {file_path}")
            return jsonify({'error': '文件不存在'}), 404
    except Exception as e:
        logger.error(f"预览文件错误: {e}")
        return jsonify({'error': f'预览文件错误: {str(e)}'}), 500

@app.route('/api/download/<filename>')
def download_file(filename):
    """下载生成的文档"""
    logger.info(f"请求下载文件: {filename}")
    
    try:
        file_path = os.path.join(tempfile.gettempdir(), filename)
        
        if os.path.exists(file_path):
            logger.info(f"文件存在，准备下载: {file_path}")
            return send_file(file_path, as_attachment=True, download_name='公文.docx')
        else:
            logger.error(f"文件不存在: {file_path}")
            return jsonify({'error': '文件不存在'}), 404
    except Exception as e:
        logger.error(f"下载文件错误: {e}")
        return jsonify({'error': f'下载文件错误: {str(e)}'}), 500

if __name__ == '__main__':
    print("启动公文生成系统后端服务...")
    print("数据库配置:", DB_CONFIG['host'])
    print("DeepSeek API配置:", "已配置" if DEEPSEEK_API_KEY != "sk-your-api-key-here" else "未配置")
    
    app.run(debug=True, host='0.0.0.0', port=5003)