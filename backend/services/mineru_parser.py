"""
MinerU文档解析服务
支持多种文档格式的智能解析
"""
import os
import tempfile
import logging
from typing import Dict, List, Any, Optional
import docx
import PyPDF2
import pandas as pd
import json
import re

logger = logging.getLogger(__name__)

class MinerUParser:
    """MinerU文档解析服务"""
    
    def __init__(self):
        """初始化解析器"""
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.csv': self._parse_csv
        }
    
    def parse_document(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """
        解析文档
        
        Args:
            file_data: 文件数据
            file_name: 文件名
            
        Returns:
            解析结果
        """
        try:
            file_ext = os.path.splitext(file_name)[1].lower()
            
            if file_ext not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'不支持的文件格式: {file_ext}'
                }
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            try:
                # 解析文档
                parse_result = self.supported_formats[file_ext](temp_file_path, file_name)
                
                # 清理临时文件
                os.unlink(temp_file_path)
                
                return parse_result
                
            except Exception as e:
                # 清理临时文件
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                raise e
                
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            return {
                'success': False,
                'error': f'文档解析失败: {str(e)}'
            }
    
    def _parse_pdf(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """解析PDF文件"""
        try:
            content = ""
            metadata = {
                'file_type': 'pdf',
                'pages': 0,
                'title': '',
                'author': '',
                'subject': ''
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata['pages'] = len(pdf_reader.pages)
                
                # 提取文档信息
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
                
                # 提取文本内容
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
            
            return {
                'success': True,
                'content': content.strip(),
                'content_length': len(content),
                'metadata': metadata,
                'file_type': 'pdf'
            }
            
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            return {
                'success': False,
                'error': f'PDF解析失败: {str(e)}'
            }
    
    def _parse_docx(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """解析Word文档"""
        try:
            content = ""
            metadata = {
                'file_type': 'docx',
                'paragraphs': 0,
                'tables': 0,
                'images': 0
            }
            
            doc = docx.Document(file_path)
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n"
                    metadata['paragraphs'] += 1
            
            # 提取表格内容
            for table in doc.tables:
                metadata['tables'] += 1
                content += "\n--- 表格 ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
            
            # 统计图片数量
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata['images'] += 1
            
            return {
                'success': True,
                'content': content.strip(),
                'content_length': len(content),
                'metadata': metadata,
                'file_type': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return {
                'success': False,
                'error': f'Word文档解析失败: {str(e)}'
            }
    
    def _parse_text(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """解析文本文件"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                return {
                    'success': False,
                    'error': '无法解析文本文件编码'
                }
            
            metadata = {
                'file_type': 'text',
                'lines': len(content.split('\n')),
                'words': len(content.split()),
                'encoding': encoding
            }
            
            return {
                'success': True,
                'content': content,
                'content_length': len(content),
                'metadata': metadata,
                'file_type': 'text'
            }
            
        except Exception as e:
            logger.error(f"文本文件解析失败: {e}")
            return {
                'success': False,
                'error': f'文本文件解析失败: {str(e)}'
            }
    
    def _parse_markdown(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # 提取标题
            titles = re.findall(r'^#{1,6}\s+(.+)$', content, re.MULTILINE)
            
            metadata = {
                'file_type': 'markdown',
                'lines': len(content.split('\n')),
                'words': len(content.split()),
                'titles': titles,
                'title_count': len(titles)
            }
            
            return {
                'success': True,
                'content': content,
                'content_length': len(content),
                'metadata': metadata,
                'file_type': 'markdown'
            }
            
        except Exception as e:
            logger.error(f"Markdown文件解析失败: {e}")
            return {
                'success': False,
                'error': f'Markdown文件解析失败: {str(e)}'
            }
    
    def _parse_excel(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """解析Excel文件"""
        try:
            content = ""
            metadata = {
                'file_type': 'excel',
                'sheets': 0,
                'total_rows': 0,
                'total_columns': 0
            }
            
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            metadata['sheets'] = len(excel_file.sheet_names)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                content += f"\n--- 工作表: {sheet_name} ---\n"
                content += f"行数: {len(df)}, 列数: {len(df.columns)}\n"
                
                # 添加列名
                content += "列名: " + " | ".join(df.columns.tolist()) + "\n"
                
                # 添加前10行数据
                for i, row in df.head(10).iterrows():
                    row_text = " | ".join([str(cell) for cell in row.values])
                    content += row_text + "\n"
                
                if len(df) > 10:
                    content += f"... (还有 {len(df) - 10} 行数据)\n"
                
                metadata['total_rows'] += len(df)
                metadata['total_columns'] = max(metadata['total_columns'], len(df.columns))
            
            return {
                'success': True,
                'content': content.strip(),
                'content_length': len(content),
                'metadata': metadata,
                'file_type': 'excel'
            }
            
        except Exception as e:
            logger.error(f"Excel文件解析失败: {e}")
            return {
                'success': False,
                'error': f'Excel文件解析失败: {str(e)}'
            }
    
    def _parse_csv(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """解析CSV文件"""
        try:
            content = ""
            metadata = {
                'file_type': 'csv',
                'rows': 0,
                'columns': 0
            }
            
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                return {
                    'success': False,
                    'error': '无法解析CSV文件编码'
                }
            
            metadata['rows'] = len(df)
            metadata['columns'] = len(df.columns)
            
            content += f"行数: {len(df)}, 列数: {len(df.columns)}\n"
            content += "列名: " + " | ".join(df.columns.tolist()) + "\n"
            
            # 添加前20行数据
            for i, row in df.head(20).iterrows():
                row_text = " | ".join([str(cell) for cell in row.values])
                content += row_text + "\n"
            
            if len(df) > 20:
                content += f"... (还有 {len(df) - 20} 行数据)\n"
            
            return {
                'success': True,
                'content': content.strip(),
                'content_length': len(content),
                'metadata': metadata,
                'file_type': 'csv'
            }
            
        except Exception as e:
            logger.error(f"CSV文件解析失败: {e}")
            return {
                'success': False,
                'error': f'CSV文件解析失败: {str(e)}'
            }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        文本分块
        
        Args:
            text: 文本内容
            chunk_size: 块大小
            overlap: 重叠大小
            
        Returns:
            文本块列表
        """
        try:
            chunks = []
            
            if len(text) <= chunk_size:
                chunks.append({
                    'content': text,
                    'start': 0,
                    'end': len(text),
                    'size': len(text)
                })
            else:
                start = 0
                while start < len(text):
                    end = min(start + chunk_size, len(text))
                    
                    # 如果不是最后一块，尝试在句号处分割
                    if end < len(text):
                        # 在块末尾附近查找句号
                        for i in range(end, max(start + chunk_size - 100, start), -1):
                            if text[i] in '。！？.!?':
                                end = i + 1
                                break
                    
                    chunk_text = text[start:end].strip()
                    if chunk_text:
                        chunks.append({
                            'content': chunk_text,
                            'start': start,
                            'end': end,
                            'size': len(chunk_text)
                        })
                    
                    # 计算下一个块的起始位置
                    start = max(start + 1, end - overlap)
            
            logger.info(f"文本分块完成，共 {len(chunks)} 块")
            return chunks
            
        except Exception as e:
            logger.error(f"文本分块失败: {e}")
            return [] 