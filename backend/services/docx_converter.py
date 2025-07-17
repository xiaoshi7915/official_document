import os
import subprocess
import uuid
import base64
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mammoth
import io
import tempfile

class DocxConverter:
    """Word文档转换服务"""
    
    def __init__(self, output_dir="temp"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
    
    def docx_to_html(self, docx_path):
        """将Word文档转换为HTML"""
        try:
            print(f"转换Word文档为HTML: {docx_path}")
            
            # 检查文件是否存在
            if not os.path.exists(docx_path):
                print(f"文件不存在: {docx_path}")
                return None, f"文件不存在: {docx_path}"
            
            # 使用mammoth库转换docx为html
            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                
                # 添加基本样式
                styled_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{
                            font-family: SimSun, "宋体", "仿宋", FangSong, serif;
                            line-height: 1.5;
                            margin: 0;
                            padding: 20px;
                        }}
                        .document {{
                            width: 100%;
                            max-width: 800px;
                            margin: 0 auto;
                            border: 1px solid #ddd;
                            padding: 40px;
                            box-shadow: 0 0 10px rgba(0,0,0,0.1);
                            background-color: white;
                        }}
                        h1, h2, h3 {{
                            text-align: center;
                            font-weight: bold;
                        }}
                        .center {{
                            text-align: center;
                        }}
                        .right {{
                            text-align: right;
                        }}
                        .red {{
                            color: #d32f2f;
                        }}
                        .bold {{
                            font-weight: bold;
                        }}
                        .indent {{
                            text-indent: 2em;
                        }}
                        table {{
                            width: 100%;
                            border-collapse: collapse;
                        }}
                        td, th {{
                            border: 1px solid #ddd;
                            padding: 8px;
                        }}
                    </style>
                </head>
                <body>
                    <div class="document">
                        {html}
                    </div>
                </body>
                </html>
                """
                
                return styled_html, None
        except Exception as e:
            print(f"转换Word文档为HTML出错: {str(e)}")
            return None, str(e)
    
    def docx_to_image(self, docx_path):
        """将Word文档转换为图片"""
        try:
            print(f"转换Word文档为图片: {docx_path}")
            
            # 检查文件是否存在
            if not os.path.exists(docx_path):
                print(f"文件不存在: {docx_path}")
                return None, f"文件不存在: {docx_path}"
            
            # 创建临时文件
            temp_dir = tempfile.mkdtemp()
            temp_html = os.path.join(temp_dir, f"{uuid.uuid4()}.html")
            temp_png = os.path.join(temp_dir, f"{uuid.uuid4()}.png")
            
            # 方法1: 使用HTML预览图片
            try:
                # 先转换为HTML
                html_content, _ = self.docx_to_html(docx_path)
                if html_content:
                    # 保存HTML到临时文件
                    with open(temp_html, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                    
                    # 使用PIL创建一个空白图片
                    from PIL import Image, ImageDraw, ImageFont
                    
                    # 创建一个带有文本的图片
                    img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
                    d = ImageDraw.Draw(img)
                    
                    # 尝试加载字体
                    try:
                        font = ImageFont.truetype("simfang.ttf", 16)
                    except:
                        try:
                            font = ImageFont.truetype("simsun.ttc", 16)
                        except:
                            font = ImageFont.load_default()
                    
                    # 提取文档结构
                    structure, _ = self.extract_docx_structure(docx_path)
                    
                    if structure and 'paragraphs' in structure:
                        # 绘制段落
                        y_position = 50
                        for para in structure['paragraphs']:
                            text = para['text']
                            alignment = para['alignment']
                            
                            # 计算文本位置
                            if alignment == 'center':
                                x_position = 400 - (len(text) * 8) / 2
                            elif alignment == 'right':
                                x_position = 700 - (len(text) * 8)
                            else:
                                x_position = 50
                            
                            # 绘制文本
                            d.text((x_position, y_position), text, fill=(0, 0, 0), font=font)
                            y_position += 30
                    else:
                        # 如果无法提取结构，则显示简单文本
                        doc = Document(docx_path)
                        y_position = 50
                        for para in doc.paragraphs:
                            if para.text.strip():
                                d.text((50, y_position), para.text, fill=(0, 0, 0), font=font)
                                y_position += 30
                    
                    # 保存图片
                    img.save(temp_png)
                    
                    # 读取图片并转换为base64
                    with open(temp_png, "rb") as image_file:
                        encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
                    
                    return encoded_string, None
            except Exception as e:
                print(f"使用PIL创建图片失败: {str(e)}")
            
            # 方法2: 使用预定义的模板图片
            try:
                # 获取模板ID
                template_id = os.path.basename(docx_path).replace('.docx', '')
                
                # 检查是否有预定义的模板图片
                template_image_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 
                                                 "frontend", "public", "template_images", f"{template_id}.png")
                
                if os.path.exists(template_image_path):
                    print(f"使用预定义的模板图片: {template_image_path}")
                    
                    # 读取图片并转换为base64
                    with open(template_image_path, "rb") as image_file:
                        encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
                    
                    return encoded_string, None
            except Exception as e:
                print(f"使用预定义的模板图片失败: {str(e)}")
            
            # 方法3: 创建一个简单的图片，显示"预览不可用"
            try:
                from PIL import Image, ImageDraw, ImageFont
                
                # 创建一个带有文本的图片
                img = Image.new('RGB', (800, 600), color=(255, 255, 255))
                d = ImageDraw.Draw(img)
                
                # 尝试加载字体
                try:
                    font = ImageFont.truetype("simfang.ttf", 24)
                except:
                    try:
                        font = ImageFont.truetype("simsun.ttc", 24)
                    except:
                        font = ImageFont.load_default()
                
                # 绘制文本
                d.text((250, 250), "预览不可用，请使用HTML预览", fill=(0, 0, 0), font=font)
                d.text((200, 300), "请确保已安装必要的依赖和工具", fill=(0, 0, 0), font=font)
                
                # 保存图片
                img.save(temp_png)
                
                # 读取图片并转换为base64
                with open(temp_png, "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
                
                return encoded_string, "无法转换Word文档为图片，使用默认图片代替"
            except Exception as e:
                print(f"创建默认图片失败: {str(e)}")
            
            # 所有方法都失败了
            return None, "无法将Word文档转换为图片，所有转换方法都失败"
        except Exception as e:
            print(f"转换Word文档为图片出错: {str(e)}")
            return None, str(e)
        finally:
            # 清理临时文件
            try:
                import shutil
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except:
                pass
    
    def extract_docx_structure(self, docx_path):
        """提取Word文档的结构信息"""
        try:
            print(f"提取Word文档结构: {docx_path}")
            
            # 检查文件是否存在
            if not os.path.exists(docx_path):
                print(f"文件不存在: {docx_path}")
                return None, f"文件不存在: {docx_path}"
            
            # 打开Word文档
            doc = Document(docx_path)
            
            # 提取段落信息
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # 获取段落对齐方式
                    alignment = "left"
                    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment = "center"
                    elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        alignment = "right"
                    elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment = "justify"
                    
                    # 获取字体信息
                    font_info = {}
                    if para.runs:
                        run = para.runs[0]
                        font_info = {
                            "name": run.font.name if run.font.name else "默认字体",
                            "size": run.font.size.pt if run.font.size else 12,
                            "bold": run.font.bold if run.font.bold is not None else False,
                            "italic": run.font.italic if run.font.italic is not None else False,
                            "color": run.font.color.rgb if run.font.color and run.font.color.rgb else None
                        }
                    
                    paragraphs.append({
                        "text": para.text,
                        "alignment": alignment,
                        "style": para.style.name if para.style else "Normal",
                        "font": font_info
                    })
            
            # 提取表格信息
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                tables.append(table_data)
            
            # 提取页面设置
            section = doc.sections[0]
            page_setup = {
                "width": section.page_width.inches,
                "height": section.page_height.inches,
                "left_margin": section.left_margin.inches,
                "right_margin": section.right_margin.inches,
                "top_margin": section.top_margin.inches,
                "bottom_margin": section.bottom_margin.inches
            }
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "page_setup": page_setup
            }, None
        except Exception as e:
            print(f"提取Word文档结构出错: {str(e)}")
            return None, str(e)