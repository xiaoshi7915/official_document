"""
文档解析服务类
实现MinerU智能解析功能，支持多种文档格式
"""
import os
import tempfile
from typing import Dict, List, Optional, Any
import io

# 文档解析库
import PyPDF2
from docx import Document
import mammoth
import pandas as pd
import openpyxl

from config_rag import SUPPORTED_FILE_TYPES

# 导入统一的日志管理器
try:
    from utils.logger import get_service_logger
    logger = get_service_logger('document_parser')
except ImportError:
    import logging
    logger = logging.getLogger(__name__)

class DocumentParser:
    """文档解析服务类"""
    
    def __init__(self):
        """初始化文档解析器"""
        self.supported_types = SUPPORTED_FILE_TYPES
        logger.info("文档解析器初始化成功")
    
    def parse_document(self, file_data: bytes, file_name: str, content_type: str = None) -> Dict[str, Any]:
        """
        解析文档内容
        
        Args:
            file_data: 文件数据
            file_name: 文件名
            content_type: 内容类型
            
        Returns:
            解析结果字典
        """
        try:
            # 获取文件扩展名
            file_ext = os.path.splitext(file_name)[1].lower().lstrip('.')
            
            # 检查文件类型是否支持
            if file_ext not in self.supported_types:
                raise ValueError(f"不支持的文件类型: {file_ext}")
            
            # 根据文件类型选择解析方法
            if file_ext == 'pdf':
                content = self._parse_pdf(file_data)
            elif file_ext in ['docx', 'doc']:
                content = self._parse_word(file_data, file_ext)
            elif file_ext == 'txt':
                content = self._parse_text(file_data)
            elif file_ext == 'md':
                content = self._parse_markdown(file_data)
            elif file_ext in ['xlsx', 'xls']:
                content = self._parse_excel(file_data, file_ext)
            elif file_ext == 'csv':
                content = self._parse_csv(file_data)
            else:
                raise ValueError(f"未实现的文件类型解析: {file_ext}")
            
            # 构建解析结果
            result = {
                'file_name': file_name,
                'file_type': file_ext,
                'content': content,
                'content_length': len(content),
                'parse_success': True,
                'error_message': None
            }
            
            logger.info(f"文档解析成功: {file_name}, 内容长度: {len(content)}")
            return result
            
        except Exception as e:
            logger.error(f"文档解析失败: {file_name}, 错误: {e}")
            return {
                'file_name': file_name,
                'file_type': file_ext if 'file_ext' in locals() else 'unknown',
                'content': '',
                'content_length': 0,
                'parse_success': False,
                'error_message': str(e)
            }
    
    def _parse_pdf(self, file_data: bytes) -> str:
        """解析PDF文件"""
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            content = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    content.append(f"第{page_num + 1}页:\n{page_text}")
            
            return '\n\n'.join(content)
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise
    
    def _parse_word(self, file_data: bytes, file_ext: str) -> str:
        """解析Word文档"""
        try:
            if file_ext == 'docx':
                # 解析docx文件
                doc = Document(io.BytesIO(file_data))
                content = []
                
                # 提取段落文本
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                
                # 提取表格文本
                for table in doc.tables:
                    table_content = []
                    for row in table.rows:
                        row_content = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_content.append(cell.text.strip())
                        if row_content:
                            table_content.append(' | '.join(row_content))
                    if table_content:
                        content.append('\n'.join(table_content))
                
                return '\n\n'.join(content)
            
            elif file_ext == 'doc':
                # 解析doc文件（使用mammoth）
                result = mammoth.extract_raw_text(io.BytesIO(file_data))
                return result.value
            else:
                raise ValueError(f"不支持的Word格式: {file_ext}")
                
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            raise
    
    def _parse_text(self, file_data: bytes) -> str:
        """解析文本文件"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    return file_data.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用latin-1
            return file_data.decode('latin-1')
            
        except Exception as e:
            logger.error(f"文本文件解析失败: {e}")
            raise
    
    def _parse_markdown(self, file_data: bytes) -> str:
        """解析Markdown文件"""
        try:
            # Markdown文件按文本处理
            return self._parse_text(file_data)
        except Exception as e:
            logger.error(f"Markdown文件解析失败: {e}")
            raise
    
    def _parse_excel(self, file_data: bytes, file_ext: str) -> str:
        """解析Excel文件"""
        try:
            content = []
            
            if file_ext == 'xlsx':
                # 读取Excel文件
                excel_file = io.BytesIO(file_data)
                excel_data = pd.read_excel(excel_file, sheet_name=None)
                
                for sheet_name, df in excel_data.items():
                    if not df.empty:
                        content.append(f"工作表: {sheet_name}")
                        # 转换为文本格式
                        content.append(df.to_string(index=False))
                        content.append("")  # 空行分隔
                
                return '\n'.join(content)
            
            elif file_ext == 'xls':
                # 读取旧版Excel文件
                excel_file = io.BytesIO(file_data)
                excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
                
                for sheet_name, df in excel_data.items():
                    if not df.empty:
                        content.append(f"工作表: {sheet_name}")
                        content.append(df.to_string(index=False))
                        content.append("")
                
                return '\n'.join(content)
            else:
                raise ValueError(f"不支持的Excel格式: {file_ext}")
                
        except Exception as e:
            logger.error(f"Excel文件解析失败: {e}")
            raise
    
    def _parse_csv(self, file_data: bytes) -> str:
        """解析CSV文件"""
        try:
            csv_file = io.BytesIO(file_data)
            
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312']
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(csv_file, encoding=encoding)
                    csv_file.seek(0)  # 重置文件指针
                    return df.to_string(index=False)
                except UnicodeDecodeError:
                    csv_file.seek(0)  # 重置文件指针
                    continue
                except Exception:
                    csv_file.seek(0)  # 重置文件指针
                    continue
            
            # 如果所有编码都失败，使用latin-1
            csv_file.seek(0)
            df = pd.read_csv(csv_file, encoding='latin-1')
            return df.to_string(index=False)
            
        except Exception as e:
            logger.error(f"CSV文件解析失败: {e}")
            raise
    
    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """
        提取文档元数据
        
        Args:
            file_data: 文件数据
            file_name: 文件名
            
        Returns:
            元数据字典
        """
        try:
            file_ext = os.path.splitext(file_name)[1].lower().lstrip('.')
            metadata = {
                'file_name': file_name,
                'file_type': file_ext,
                'file_size': len(file_data),
                'has_content': False
            }
            
            # 根据文件类型提取特定元数据
            if file_ext == 'pdf':
                pdf_metadata = self._extract_pdf_metadata(file_data)
                metadata.update(pdf_metadata)
            elif file_ext in ['docx', 'doc']:
                word_metadata = self._extract_word_metadata(file_data, file_ext)
                metadata.update(word_metadata)
            elif file_ext in ['xlsx', 'xls']:
                excel_metadata = self._extract_excel_metadata(file_data, file_ext)
                metadata.update(excel_metadata)
            
            return metadata
            
        except Exception as e:
            logger.error(f"元数据提取失败: {e}")
            return {
                'file_name': file_name,
                'file_type': file_ext if 'file_ext' in locals() else 'unknown',
                'file_size': len(file_data),
                'has_content': False,
                'error': str(e)
            }
    
    def _extract_pdf_metadata(self, file_data: bytes) -> Dict[str, Any]:
        """提取PDF元数据"""
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            metadata = {
                'page_count': len(pdf_reader.pages),
                'has_content': len(pdf_reader.pages) > 0
            }
            
            # 提取PDF信息
            if pdf_reader.metadata:
                info = pdf_reader.metadata
                if info.get('/Title'):
                    metadata['title'] = info['/Title']
                if info.get('/Author'):
                    metadata['author'] = info['/Author']
                if info.get('/Subject'):
                    metadata['subject'] = info['/Subject']
                if info.get('/Creator'):
                    metadata['creator'] = info['/Creator']
            
            return metadata
        except Exception as e:
            logger.error(f"PDF元数据提取失败: {e}")
            return {'has_content': False}
    
    def _extract_word_metadata(self, file_data: bytes, file_ext: str) -> Dict[str, Any]:
        """提取Word文档元数据"""
        try:
            if file_ext == 'docx':
                doc = Document(io.BytesIO(file_data))
                
                metadata = {
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'has_content': len(doc.paragraphs) > 0 or len(doc.tables) > 0
                }
                
                # 提取核心属性
                core_props = doc.core_properties
                if core_props.title:
                    metadata['title'] = core_props.title
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.subject:
                    metadata['subject'] = core_props.subject
                if core_props.created:
                    metadata['created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['modified'] = core_props.modified.isoformat()
                
                return metadata
            else:
                return {'has_content': True}  # doc文件默认有内容
                
        except Exception as e:
            logger.error(f"Word元数据提取失败: {e}")
            return {'has_content': False}
    
    def _extract_excel_metadata(self, file_data: bytes, file_ext: str) -> Dict[str, Any]:
        """提取Excel文件元数据"""
        try:
            excel_file = io.BytesIO(file_data)
            
            if file_ext == 'xlsx':
                workbook = openpyxl.load_workbook(excel_file, read_only=True)
                metadata = {
                    'sheet_count': len(workbook.sheetnames),
                    'sheet_names': workbook.sheetnames,
                    'has_content': len(workbook.sheetnames) > 0
                }
            else:
                # 对于xls文件，使用pandas获取基本信息
                excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
                metadata = {
                    'sheet_count': len(excel_data),
                    'sheet_names': list(excel_data.keys()),
                    'has_content': len(excel_data) > 0
                }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Excel元数据提取失败: {e}")
            return {'has_content': False}
    
    def validate_file(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """
        验证文件
        
        Args:
            file_data: 文件数据
            file_name: 文件名
            
        Returns:
            验证结果
        """
        try:
            file_ext = os.path.splitext(file_name)[1].lower().lstrip('.')
            
            # 检查文件类型
            if file_ext not in self.supported_types:
                return {
                    'valid': False,
                    'error': f"不支持的文件类型: {file_ext}"
                }
            
            # 检查文件大小
            file_size = len(file_data)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                return {
                    'valid': False,
                    'error': f"文件大小超过限制: {file_size / 1024 / 1024:.2f}MB > 50MB"
                }
            
            # 检查文件内容
            if file_size == 0:
                return {
                    'valid': False,
                    'error': "文件为空"
                }
            
            # 提取元数据验证文件完整性
            metadata = self.extract_metadata(file_data, file_name)
            if not metadata.get('has_content', False):
                return {
                    'valid': False,
                    'error': "文件内容为空或无法解析"
                }
            
            return {
                'valid': True,
                'file_type': file_ext,
                'file_size': file_size,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"文件验证失败: {e}")
            return {
                'valid': False,
                'error': f"文件验证失败: {str(e)}"
            } 