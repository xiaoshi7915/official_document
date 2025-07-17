import os
import markdown
import docx2txt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import uuid
from typing import Dict, Any

class DocumentProcessor:
    """文档处理服务"""
    
    def __init__(self):
        self.output_dir = "output"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def parse_uploaded_file(self, file_path: str) -> str:
        """解析上传的文件内容"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext in ['.docx', '.doc']:
            return docx2txt.process(file_path)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def generate_document(self, content: str, template_type: str, metadata: Dict[str, Any]) -> str:
        """生成公文文档"""
        try:
            print(f"开始生成文档，模板类型: {template_type}")
            print(f"元数据: {metadata}")
            
            # 创建新文档
            doc = Document()
            
            # 设置页面格式
            self._setup_page_format(doc)
            
            # 添加文档头部
            self._add_document_header(doc, template_type, metadata)
            
            # 处理内容
            format_type = metadata.get('format_type', 'plain')
            print(f"内容格式: {format_type}")
            
            if format_type == 'markdown':
                self._add_markdown_content(doc, content)
            else:
                self._add_plain_content(doc, content)
                
            # 添加文档尾部
            self._add_document_footer(doc, metadata)
            
            # 保存文档
            filename = f"{template_type}_{uuid.uuid4().hex[:8]}.docx"
            output_path = os.path.join(self.output_dir, filename)
            
            print(f"保存文档到: {output_path}")
            doc.save(output_path)
            
            print("文档生成完成")
            return output_path
        except Exception as e:
            print(f"生成文档时出错: {str(e)}")
            raise
    
    def _setup_page_format(self, doc: Document):
        """设置页面格式"""
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4纸高度
        section.page_width = Inches(8.27)    # A4纸宽度
        section.left_margin = Inches(1.18)   # 左边距30mm
        section.right_margin = Inches(0.79)  # 右边距20mm
        section.top_margin = Inches(1.46)    # 上边距37mm
        section.bottom_margin = Inches(1.18) # 下边距30mm
    
    def _add_document_header(self, doc: Document, template_type: str, metadata: Dict[str, Any]):
        """添加公文头部"""
        # 发文机关标识
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run(metadata.get('sender', '发文机关'))
        header_run.font.name = '方正小标宋简体'
        header_run.font.size = Pt(22)
        header_run.bold = True
        
        # 分隔线
        line_p = doc.add_paragraph()
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_p.add_run('━' * 30)
        line_run.font.name = '仿宋'
        line_run.font.size = Pt(14)
        
        # 标题
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(metadata.get('title', '标题'))
        title_run.font.name = '方正小标宋简体'
        title_run.font.size = Pt(22)
        title_run.bold = True
        
        # 文号等信息
        if metadata.get('document_number'):
            doc_num_p = doc.add_paragraph()
            doc_num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_num_run = doc_num_p.add_run(metadata['document_number'])
            doc_num_run.font.name = '仿宋'
            doc_num_run.font.size = Pt(16)
        
        # 空行
        doc.add_paragraph()
    
    def _add_markdown_content(self, doc: Document, content: str):
        """添加Markdown格式的内容"""
        # 简单的Markdown解析
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            p = doc.add_paragraph()
            
            # 处理标题
            if line.startswith('# '):
                run = p.add_run(line[2:])
                run.font.name = '黑体'
                run.font.size = Pt(16)
                run.bold = True
            elif line.startswith('## '):
                run = p.add_run(line[3:])
                run.font.name = '黑体'
                run.font.size = Pt(14)
                run.bold = True
            else:
                run = p.add_run(line)
                run.font.name = '仿宋'
                run.font.size = Pt(16)
            
            # 设置行距
            p.paragraph_format.line_spacing = 1.5
    
    def _add_plain_content(self, doc: Document, content: str):
        """添加纯文本内容"""
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(para_text.strip())
                run.font.name = '仿宋'
                run.font.size = Pt(16)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Inches(0.39)  # 首行缩进2字符
    
    def _add_document_footer(self, doc: Document, metadata: Dict[str, Any]):
        """添加公文尾部"""
        # 空行
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 发文机关署名
        if metadata.get('sender'):
            sender_p = doc.add_paragraph()
            sender_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sender_run = sender_p.add_run(metadata['sender'])
            sender_run.font.name = '仿宋'
            sender_run.font.size = Pt(16)
        
        # 发文日期
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_text = metadata.get('date', datetime.now().strftime('%Y年%m月%d日'))
        date_run = date_p.add_run(date_text)
        date_run.font.name = '仿宋'
        date_run.font.size = Pt(16)