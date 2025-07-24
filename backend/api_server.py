from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import mysql.connector
import os
import json
import requests
import tempfile
from datetime import datetime
from werkzeug.utils import secure_filename
import docx
from docx import Document
from docx.shared import Inches
import mammoth

app = Flask(__name__)
CORS(app)

# 数据库配置
DB_CONFIG = {
    'host': '47.118.250.53',
    'port': 3306,
    'database': 'official_doc',
    'user': 'official_doc',
    'password': 'admin123456!'
}

# DeepSeek API配置
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-your-api-key-here")
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# 上传文件配置
UPLOAD_FOLDER = 'temp'
ALLOWED_EXTENSIONS = {'txt', 'md', 'docx'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def get_db_connection():
    """获取数据库连接"""
    try:
        return mysql.connector.connect(**DB_CONFIG)
    except mysql.connector.Error as err:
        print(f"数据库连接错误: {err}")
        return None

def init_database():
    """初始化数据库表"""
    conn = get_db_connection()
    if not conn:
        return False
    
    try:
        cursor = conn.cursor()
        
        # 创建公文类型表
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS document_types (
            id VARCHAR(50) PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            description TEXT
        )
        """)
        
        # 创建公文字段表
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS document_fields (
            id INT AUTO_INCREMENT PRIMARY KEY,
            field_name VARCHAR(100) NOT NULL,
            field_key VARCHAR(100) NOT NULL,
            field_type VARCHAR(50) NOT NULL,
            field_category VARCHAR(50) NOT NULL,
            is_required BOOLEAN DEFAULT FALSE,
            parent_field_id INT,
            document_type_id VARCHAR(50),
            FOREIGN KEY (parent_field_id) REFERENCES document_fields(id),
            FOREIGN KEY (document_type_id) REFERENCES document_types(id)
        )
        """)
        
        # 创建用户生成的公文表
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS generated_documents (
            id INT AUTO_INCREMENT PRIMARY KEY,
            document_type_id VARCHAR(50) NOT NULL,
            title VARCHAR(200) NOT NULL,
            content TEXT,
            metadata JSON,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            FOREIGN KEY (document_type_id) REFERENCES document_types(id)
        )
        """)
        
        # 插入15种公文类型
        document_types = [
            ('baogao', '报告', '向上级机关汇报工作、反映情况、回复询问'),
            ('gonggao', '公告', '向国内外宣布重要事项或者法定事项'),
            ('gongbao', '公报', '公开发布重要决议、决定或重大事件'),
            ('hansong', '函送', '向有关单位送交公文或资料'),
            ('jiyao', '纪要', '记载会议主要情况和议定事项'),
            ('jueding', '决定', '对重要事项或重大行动作出安排'),
            ('jueyi', '决议', '会议讨论通过的重要事项的决策'),
            ('mingling', '命令', '依照有关法律公布行政法规和规章、宣布施行重大强制性措施'),
            ('pifu', '批复', '答复下级机关请示事项'),
            ('qingshi', '请示', '向上级机关请求指示或批准'),
            ('tongbao', '通报', '表彰先进、批评错误、传达重要精神或情况'),
            ('tongzhi', '通知', '发布、传达要求下级机关执行和有关单位周知或者执行的事项'),
            ('tonggao', '通告', '公开宣布重要事项或者法定事项'),
            ('yian', '议案', '正式提出审议事项的文书'),
            ('yijian', '意见', '对重要问题提出见解和处理办法')
        ]
        
        for doc_type in document_types:
            cursor.execute("""
            INSERT IGNORE INTO document_types (id, name, description) 
            VALUES (%s, %s, %s)
            """, doc_type)
        
        conn.commit()
        cursor.close()
        conn.close()
        return True
    except Exception as e:
        print(f"数据库初始化错误: {e}")
        return False

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/templates', methods=['GET'])
def get_templates():
    """获取公文类型列表"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': '数据库连接失败'}), 500
    
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM document_types ORDER BY id")
    templates = cursor.fetchall()
    cursor.close()
    conn.close()
    
    return jsonify({'data': templates})

@app.route('/api/document-fields/<type_id>', methods=['GET'])
def get_document_fields(type_id):
    """获取公文字段"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': '数据库连接失败'}), 500
    
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
    SELECT * FROM document_fields 
    WHERE document_type_id = %s 
    ORDER BY parent_field_id IS NULL DESC, parent_field_id, field_category
    """, (type_id,))
    fields = cursor.fetchall()
    cursor.close()
    conn.close()
    
    return jsonify({'data': fields})

@app.route('/api/generate-title', methods=['POST'])
def generate_title():
    """从内容生成标题"""
    data = request.json
    content = data.get('content', '')
    
    if not content:
        return jsonify({'success': False, 'message': '内容不能为空'})
    
    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        json_data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": "你是一个专业的公文标题生成助手，请根据提供的公文内容生成一个简洁、准确、符合公文规范的标题，不超过20个字。"
                },
                {
                    "role": "user", 
                    "content": f"请根据以下公文内容生成标题：\n\n{content}"
                }
            ]
        }
        
        response = requests.post(
            DEEPSEEK_API_URL,
            headers=headers,
            json=json_data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            title = result['choices'][0]['message']['content'].strip()
            return jsonify({'success': True, 'title': title})
        else:
            return jsonify({'success': False, 'message': 'API调用失败'})
            
    except Exception as e:
        print(f"生成标题错误: {e}")
        return jsonify({'success': False, 'message': '生成标题失败，请稍后再试'})

@app.route('/api/generate-content', methods=['POST'])
def generate_content():
    """从主题生成内容"""
    data = request.json
    topic = data.get('topic', '')
    document_type = data.get('document_type', '')
    
    if not topic:
        return jsonify({'success': False, 'message': '主题不能为空'})
    
    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        json_data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": f"你是一个专业的公文写作助手，请根据提供的主题生成一篇符合{document_type}格式规范的公文内容。内容应该结构清晰、语言规范、符合公文写作要求。"
                },
                {
                    "role": "user",
                    "content": f"请根据以下主题生成{document_type}内容：\n\n{topic}"
                }
            ]
        }
        
        response = requests.post(
            DEEPSEEK_API_URL,
            headers=headers,
            json=json_data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content'].strip()
            return jsonify({'success': True, 'content': content})
        else:
            return jsonify({'success': False, 'message': 'API调用失败'})
            
    except Exception as e:
        print(f"生成内容错误: {e}")
        return jsonify({'success': False, 'message': '生成内容失败，请稍后再试'})

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文件并解析内容"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '没有文件'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '没有选择文件'})
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        try:
            content = ""
            
            if filename.endswith('.md'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif filename.endswith('.docx'):
                doc = docx.Document(file_path)
                content = '\n'.join([para.text for para in doc.paragraphs])
            
            # 删除临时文件
            if os.path.exists(file_path):
                os.remove(file_path)
                
            return jsonify({'success': True, 'content': content})
            
        except Exception as e:
            print(f"文件解析错误: {e}")
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'success': False, 'message': '文件解析失败'})
    
    return jsonify({'success': False, 'message': '不支持的文件格式'})

@app.route('/api/generate', methods=['POST'])
def generate_document():
    """生成公文"""
    data = request.json
    
    # 保存到数据库
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("""
            INSERT INTO generated_documents 
            (document_type_id, title, content, metadata) 
            VALUES (%s, %s, %s, %s)
            """, (
                data.get('template_type', ''),
                data.get('metadata', {}).get('title', ''),
                data.get('content', ''),
                json.dumps(data.get('metadata', {}))
            ))
            document_id = cursor.lastrowid
            conn.commit()
            cursor.close()
        except Exception as e:
            print(f"数据库保存错误: {e}")
            document_id = None
        finally:
            conn.close()
    
    # 生成Word文档
    try:
        doc = Document()
        
        # 添加标题
        title = data.get('metadata', {}).get('title', '公文标题')
        title_para = doc.add_heading(title, 0)
        title_para.alignment = 1  # 居中对齐
        
        # 添加发文机关
        sender = data.get('metadata', {}).get('sender', '')
        if sender:
            sender_para = doc.add_paragraph(sender)
            sender_para.alignment = 1  # 居中对齐
        
        # 添加文号
        doc_number = data.get('metadata', {}).get('document_number', '')
        if doc_number:
            doc_number_para = doc.add_paragraph(doc_number)
            doc_number_para.alignment = 1  # 居中对齐
        
        # 添加正文内容
        content = data.get('content', '')
        if data.get('metadata', {}).get('format_type') == 'markdown':
            # 如果是Markdown格式，简单处理
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            # 纯文本格式
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
        
        # 添加日期
        date = data.get('metadata', {}).get('date')
        if date:
            date_para = doc.add_paragraph(date)
            date_para.alignment = 2  # 右对齐
        
        # 保存文档
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return jsonify({
            'success': True,
            'document_id': document_id,
            'download_url': f'/api/download/{os.path.basename(temp_file.name)}'
        })
        
    except Exception as e:
        print(f"文档生成错误: {e}")
        return jsonify({'success': False, 'message': '文档生成失败'})

@app.route('/api/download/<filename>')
def download_file(filename):
    """下载生成的文档"""
    file_path = os.path.join(tempfile.gettempdir(), filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True, download_name='公文.docx')
    else:
        return jsonify({'error': '文件不存在'}), 404

if __name__ == '__main__':
    # 初始化数据库
    if init_database():
        print("数据库初始化成功")
    else:
        print("数据库初始化失败")
    
    app.run(host='0.0.0.0', port=5003, debug=True)